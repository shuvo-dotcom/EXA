# -*- coding: utf-8 -*-
"""
Created on Tue Apr 30 00:25:36 2024
@author: ENTSOE
"""

import sys
import os
import pandas as pd 
import warnings
import re 
import time
import json
import concurrent.futures
from fpdf import FPDF
from docx import Document
import time 
import yaml

sys.path.append('utils')
sys.path.append('functions/plexos_functions')
sys.path.append('functions')
sys.path.append('functions/LLMs')

# Add import for os and ensure project root is in sys.path for correct module resolution
top_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir))
if top_dir not in sys.path:
    sys.path.insert(0, top_dir)

import src.LOLA.extract_excel_sheets as eps
from src.ai.llm_calls.open_ai_calls import run_open_ai_ns as oaicw
from src.ai.llm_calls.open_ai_calls import open_ai_search
# from tts import tts

# Dictionary of EU28 countries with their neighboring countries

def create_pdf_from_final_report(final_report, output_path="Joule_Report.pdf"):
    print("Creating PDF from final report...")
    pdf = FPDF()
    pdf.set_font("Arial", size=12)
    pdf.add_page()
    for _, row in final_report.iterrows():
        pdf.set_font("Arial", "B", 16, uni=True)  # H1
        pdf.multi_cell(0, 10, str(row.get("objective_id", "")))

        pdf.set_font("Arial", "B", 14, uni=True)  # H2
        pdf.multi_cell(0, 10, str(row.get("task_id", "")))

        pdf.set_font("Arial", "B", 12, uni=True)  # H3
        pdf.multi_cell(0, 10, str(row.get("sub_task_id", "")))

        pdf.set_font("Arial", "B", 10, uni=True)  # H4
        pdf.multi_cell(0, 10, str(row.get("country", "")))

        pdf.set_font("Arial", "", 11, uni=True)   # Paragraph
        pdf.multi_cell(0, 10, str(row.get("final_copy", "")))

    try:
        pdf.output(output_path)
        print("PDF created successfully!")
        os.startfile(output_path)
    except Exception as e:
        print("Error creating PDF:", e)

def save_file(file_path, doc, file_type="word"):
    if file_type == "word":
        while True:
            try:
                doc.save(file_path)
                break
            except PermissionError:
                print("Please close the file to continue")
                tts("Please close the file to continue")
                file_close = input("Is the file closed? (yes/no): ")
                if file_close.lower() == "yes":
                    break
                time.sleep(5)

    if file_type == 'csv':
        while True:
            try:
                doc.to_csv(file_path, index=False)
                break
            except PermissionError:
                print("Please close the file to continue")
                tts("Please close the file to continue")
                file_close = input("Is the file closed? (yes/no): ")
                if file_close.lower() == "yes":
                    break
                time.sleep(5)

def feedback_sense_check(report_extract_dict, context, main_copy):
    objective_title = report_extract_dict['objective_title']
    task_title = report_extract_dict['task_title']
    sub_task_title = report_extract_dict['sub_task_title']
    country = report_extract_dict['country']
    final_copy = report_extract_dict['final_copy']

    context = f"""
                {main_context}.
                You are asked to sense check the following section of the report. The year is 2050 so the energy landscape will have changed compared to today, 
                but for example, it's unlikely that all regions in a country have the exact same demand or energy mix, it mean boundary conditions have been reached in the simulation and may to too limited,
                or haven't been thought through thoroughly enough if all regions are the same.
                Feedback here could be that data should vary according to the region, e.g. population, landscape for building new production sources, etc.
                """

    prompt = f"""
            {context}
            You are asked to sense check the following section of the report:
            Objective: {objective_title}
            Task: {task_title}
            Sub-task: {sub_task_title}
            Country: {country}
            Information: {main_copy}
            """
    feedback = oaicw(prompt, context, model=sota_model)
    return feedback

def create_word_file(final_report, output_path):
    print("Creating Word file from final report...")
    doc = Document()
    #try to open the file if it's opened then give me an input that i can say yes to when i close the file then continue

    for obj_id, obj_group in final_report.groupby("objective_id"):
        obj_title = obj_group["objective_title"].iloc[0]
        doc.add_heading(f"{obj_id} {obj_title}", level=1)

        for t_id, task_group in obj_group.groupby("task_id"):
            t_title = task_group["task_title"].iloc[0]
            if t_title != obj_title:
                doc.add_heading(f"{t_id} {t_title}", level=2)

            for s_id, sub_task_group in task_group.groupby("sub_task_id"):
                s_title = sub_task_group["sub_task_title"].iloc[0]
                if s_title != t_title:
                    doc.add_heading(f"{s_id} {s_title}", level=3)

                for country, country_group in sub_task_group.groupby("country"):
                    if country != "EU" and country != "summary" and country != s_title:
                        doc.add_heading(country, level=4)

                    for _, row in country_group.iterrows():
                        final_copy = str(row.get("final_copy", ""))
                        doc.add_paragraph(final_copy)

    save_file(output_path, doc, doc_type="word")
    print("Word file created successfully!")

def create_heating(context, heading_level):
    prompt = (
                f'Write a header Level: {heading_level}. Context: {context}. Please suggest a header name'
                'Heading Name:')
    # print(prompt)
    return prompt

def data_analysis_report(instruction, json_data_output, chapter):
    prompt = (f"""{main_context}. 
                You are currently drafting chapter {chapter}.
                Instruction: {instruction}.
                Do not end the chapter with any conclusion, unless it is stated we are at the conclusions chapter.
                Do not include a header in the response, the text will be used directly directed to in the context.
                Do neutral language, do not overuse adjectives. Use simple language.
                """)
    return prompt

def extract_response_info(first_draft):
    # input_tokens = re.search(r"input tokens: (\d+)", first_draft)
    # output_tokens = re.search(r"output tokens used: (\d+)", first_draft)
    
    graph_type = re.search(r"Graph for figure: (\w+ \w+)", first_draft)
    body_text = re.search(r"Body of text:\n(.+)", first_draft, re.DOTALL)
    
    # Extract and print the information
    # input_tokens = int(input_tokens.group(1)) if input_tokens else None
    # output_tokens = int(output_tokens.group(1)) if output_tokens else None
    graph_type = graph_type.group(1) if graph_type else None
    body_text = body_text.group(1).strip() if body_text else None
        
    return graph_type, body_text

def internet_search(country, kpi, title, description):
    context = f"""
                {main_context}.
                You are asked to analyse some additional data.
                you muust perform an internet search to obtain the data. focus on the year 2024.
                The KPI is {kpi}. The title is {title}. The description is {description}.
                
                """
    prompt = f""" 
                Please perform an internet search to obtain the data.
                The KPI is {kpi}. The title is {title}. The description is {description}. Focus on the country {country}. 
                """
    
    if 'gpt' in search_model :
        search_results = open_ai_search(prompt, model = search_model)
    if 'sonos' in search_model:
        search_results = oaicw(prompt, model = search_model)


    return search_results

def extract_response_info(first_draft):
    graph_info_regex = r"Output format - \[Graph for figure: (.*?), caption for figure: (.*?), x-axis name: (.*?), y-axis name: (.*?): Body of text:(.*?)\]"
    
    matches = re.search(graph_info_regex, first_draft, re.DOTALL)
    
    if matches:
        graph_type = matches.group(1).strip()
        caption = matches.group(2).strip()
        x_axis = matches.group(3).strip()
        y_axis = matches.group(4).strip()
        body_text = matches.group(5).strip()
        
        return graph_type, caption, x_axis, y_axis, body_text
    else:
        return None

def get_ranking(df, country, value_col, country_col ):
    df = df.groupby([country_col]).sum().reset_index()
    df['rank'] = df[value_col].rank(ascending = False)
    #get country position in ranking
    country_rank = df[df[country_col] == country]['rank'].values[0]
    country_rank_text = f'{countries[country]} is ranked {int(country_rank)} out of {len(df)} countries'
    return country_rank_text

def extract_csv_data(collection,  interval_name,  category = None, property = None):
    if interval_name.lower() == 'yearly': 
        location = os.path.join(base_location, interval_name.lower(), model_name, f'{collection.lower()}.csv')
        data_extract = pd.read_csv(location)

        if category != '-':
            try:
                valid_categories = [cat.strip() for cat in category.split(',')]
                data_extract = data_extract[data_extract['category_name'].isin(valid_categories)]
            except:
                if category != None: data_extract = data_extract[data_extract['category_name'] == category]  
                if property != None: data_extract = data_extract[data_extract['property_name'] == property]

        data_extract = data_extract[['collection_name', 'category_name', 'child_name', 'property_name', 'value', 'unit_name', 'interval_id']]

    if interval_name == 'daily':
        if collection == 'systemgasstorages':
            pasue = True

        data_extract = pd.DataFrame()
        try:
            categories = [cat.strip() for cat in category.split(',')]
            for cat in categories:
                location = os.path.join(base_location, interval_name.lower(), model_name, 
                                      f'{collection.lower()}_{cat.lower()}_{property.lower()}.csv')
                all_data = pd.read_csv(location)
                data_extract = pd.concat([data_extract, all_data], ignore_index=True)
                
                # Convert date_string to datetime and extract month-year
                monthly_data_extract = data_extract.copy()
                monthly_data_extract['month_year'] = pd.to_datetime(monthly_data_extract['date_string'], dayfirst=True, format="%d/%m/%Y").dt.to_period('M')
                
                # Group by month and other columns, excluding 'value' and 'date_string'
                group_cols = [col for col in monthly_data_extract.columns if col not in ['value', 'date_string','interval_id','period_value']]
                monthly_data_extract = monthly_data_extract.groupby(group_cols).sum().reset_index()
                monthly_data_extract.drop(columns = ['interval_id','period_value','date_string'], inplace = True)
                data_extract = monthly_data_extract 
        except Exception as e:
            location = os.path.join(base_location, interval_name.lower(), model_name, 
                                  f'{collection.lower()}_{category.lower()}_{property.lower()}.csv')
            data_extract = pd.read_csv(location)
    return data_extract

def extract_plexos_data(prompt_sheet, sheet, sub_task_id):
    if sub_task_id == '4.6.1':
        pause = True

    collection_name = prompt_sheet['collection_name']
    property_list = prompt_sheet['property_list']
    category_list = prompt_sheet['category_list']
    interval_name = prompt_sheet['Period_type']

    plexos_extract = extract_csv_data(collection_name,  interval_name, category = category_list, property = property_list)
    plexos_extract = plexos_extract[plexos_extract['value'] != 0]
    plexos_extract['country'] = plexos_extract['child_name'].str[0:2]
    return plexos_extract

def refine_and_proofread(sub_task_id, current_text, prompt_sheet, country, section_type='main'):
    """
    Refines and proofreads text, ensuring it stays concise and on-topic.
    """
    try:
        reference_text = prompt_sheet['Cross Reference'].loc[sub_task_id, 'Copy']
    except KeyError:
        reference_text = None

    refinement_prompt = f"""
                                Follow these guidelines strictly:

                                1. Maintain a consistent tone aligned with: {tone_and_feel}.
                                2. Use varied sentence structures; avoid fluff, marketing speak, or repetitive phrasing.
                                3. Proofread thoroughly for grammar, spelling, and clarity.
                                4. Write in prose suitable for a report; do not add formatting such as bold, italics, or underlines.
                                5. Since this is not an introductory section, avoid phrases like "this report," "this section," or "this subsection."
                                6. Craft the text to mimic human writing that would pass an AI detection test.

                                Authorship guidelines to adhere to:
                                - Statistical predictability (perplexity): Ensure the text includes spikes of randomness and surprise, avoiding overly smooth or predictable AI-like patterns.
                                - Uniformity of structure: Vary sentence and paragraph lengths; avoid consistent rhythm or pacing.
                                - Lexical fingerprints: Do not overuse AI-favored tokens or stock phrases (e.g., “delve,” “landscape,” “pivotal,” “robust framework”); avoid repetition across contexts.
                                - Probability curvature anomalies: Write in a way that deviates from typical AI probability basins.
                                - Watermarking signals: If applicable, avoid hidden token patterns associated with AI models.

                                Current text to refine:
                                {current_text}

                                Reference style example:
                                {reference_text}

                                Country: {country}
                                """

    context = f"""
                You are a professional editor wirting the style of IRENA and the International Energy Agency. You are a historian writing a report on the energy landscape in 2050.
                Never include the chapter name or heading in the response, this will destroy the flow of the whole report and make the output completely unusable, if they are already included in the text, remove them.
                Write in prose, do not use bullet points or lists. Do not use any formatting such as bold, italics or underlines.
                Do not create paragraph breaks, keep the text as one or 2 continuous paragraphs.
                You must ensure the text is engaging, whilst remaining professional and informative, as it is a technical report, which will be published.
                Avoid perfect grammer, use simple language, avoid jargon, and keep the language simple and natural. Throw in a very minor error to make it sound more human.
                Don't use too many adjectives.
                Add some contractions.
                Avoid repetative phrases. 
                Avoid Generic Openings such as "In today’s world...", "It is important to...", "In Chapter x we discuss", "This report will cover...", "This analysis". be more human 
                Keep language simple, natural and authoritative.
                Finally return ONLY the final copy as it will be used verbatum in the report, do not explain what adjustments you have made
    """

    # refined_text = run_claude_call(refinement_prompt, context, model=writer_model)
    try:
        refined_text = oaicw(refinement_prompt, context, model = writer_model)
    except Exception as e:
        print(f"Error in refining and proofreading text for {sub_task_id}", e), "Running with o3-mini"
        refined_text = oaicw(refinement_prompt, context, model = 'o3-mini')

    refined_text = refined_text.replace('—', ', ')
    return refined_text

def process_plexos_data(country, plexos_extract, description, prompt_sheet, row_extract):
    print('Analyzing PLEXOS data for', country)
    property = row_extract['property_list']
    categories = row_extract['category_list']
    sub_task_header = row_extract['sub_task_header']

    try:
        additional_information = row_extract['additional information']
    except:
        additional_information = None
    plexos_extract_df = plexos_extract[plexos_extract['property_name'] == property]
    
    if categories != '-':
        try:
            valid_categories = [cat.strip() for cat in categories.split(',')]
            plexos_extract_df = plexos_extract_df[plexos_extract_df['category_name'].isin(valid_categories)]
        except:
            if categories != None:
                plexos_extract_df = plexos_extract_df[plexos_extract_df['category_name'] == categories]

    category_df = prompt_sheet['All Categories'][['category_name', 'group', 'Description']]
    plexos_extract_df = plexos_extract_df.merge(category_df, on='category_name', how='left')

    categories_list = [cat.strip() for cat in categories.split(',')]
    category_descriptions = category_df[category_df['category_name'].isin(categories_list)]
    category_descriptions.drop(columns = ['category_name'], inplace = True)

    if 'Grid Flows' in sub_task_header:
        bordering_countries = country_borders[country]
        bordering_countries.append(country)
        plexos_extract_country = plexos_extract_df[plexos_extract_df['country'].isin(bordering_countries)]
    else:
        plexos_extract_country = plexos_extract_df[plexos_extract_df['country'] == country]
    
    plexos_extract_country['region'] = plexos_extract_country['child_name'].str[2:4]
    plexos_extract_country['region'] = plexos_extract_country['region'].apply(lambda x: 'Entire Country' if x == '00' else x.lstrip('0'))

    units = plexos_extract_country['unit_name'].unique()
    asset_classes = plexos_extract_country['collection_name'].unique()

    try:
        plexos_extract_country.drop(columns = ['category_name','Description', 'interval_id', 'country','property_name','unit_name'], inplace = True) 
    except:
        plexos_extract_country.drop(columns = ['model_name', 'collection_name','category_name','Description', 'country','property_name','unit_name'], inplace = True) 

    plexos_extract_country.rename(columns={'group': 'category'}, inplace=True)
    plexos_extract_country = plexos_extract_country.drop_duplicates(keep='first')
    # print(plexos_extract_country)

    try:
        plexos_extract_country = plexos_extract_country.groupby(['child_name','category', 'region']).sum().reset_index()
    except:
        pass
    #remove duplicates

    try:
        json_data_output = plexos_extract_country.to_json(orient='index')
    except:
        json_data_output = plexos_extract_country.to_dict(orient='records')

    default_data_context = data_analysis_report(description, json_data_output, description)
    
    try:
        country_rank = get_ranking(plexos_extract_df, country, value_col='value', country_col='country')
    except:
        country_rank = None

    city_names = nodalsplit[nodalsplit['Country'] == country]
    city_names['region'] = city_names['Node'].str[2:4]
    city_names['region'] = city_names['region'].apply(lambda x: 'Entire Country' if x == '00' else x.lstrip('0'))
    city_names = city_names.drop(columns = ['Node']) 
    city_names_json = city_names.to_json(orient='index')

    default_data_prompt = f"""Please give a concise summary of the data extracted. Give a neutral tone and use simple non emotional language.
                            {additional_information}.
                            Here are the city names that correspond with the regions: {city_names_json}. Consider the geographical location e.g. north, east, south west etc.
                            Start by giving a general overview of the regions, their cities and general location. If the names are too long or complex use the simple version e.g. Region de Bruxelles-Capitale = Brussels.
                            Here are the results of the simulation: {json_data_output}.
                            Here are the units of measurement: {units}.
                            Here is the country: {country}. If you are analysing grid flows you will recieve data for bordering countries as well.
                            Here are the properties used in the analysis: {property}.
                            Here are the categories used in the analysis: {categories}.
                            Here are the asset classes used in the analysis: {asset_classes}.
                            Here is where the project ranks for the data compared to other european countries: {country_rank}.
                            do not perform any formatting e.g. bold, italics.
                            """
    if plexos_extract_country.empty:
        plexos_results = None
        print('No data extracted for', country)
    else:
        plexos_results = oaicw(default_data_prompt, default_data_context, model = sota_model)
        # plexos_results = 'Test'
    return plexos_results

def extract_etm(sub_task_id, detailed_context, df, country, instruction):
    df.fillna('-', inplace=True)
    etm_extract = etm[etm['VALUE'] != 0]
    etm_extract = etm_extract.drop(columns = ['PROPERTY', 'UNIT'])
    etm_analysis = {}

    etm_extract = etm_extract[etm_extract['COUNTRY'] == country]
    sectors = etm_extract['SECTOR'].unique()

    if run_concurrent:
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future_to_sector = {
                executor.submit(
                    lambda s: {
                        s: {
                            'sector': s,
                            'final_copy': oaicw(
                                f"""Please give a concise summary of the data extracted. 
                                country: {country}. 
                                sector: {s}.
                                data: {etm_extract[etm_extract['SECTOR'] == s].drop(columns=['SECTOR'])[etm_extract['YEAR'] == 2050].rename(columns={'VALUE': 'VALUE (TWh)'}).to_json(orient='index')}
                                Compare the data between 2019 and 2050.
                                """,
                                data_analysis_report(instruction, etm_extract[etm_extract['SECTOR'] == s].drop(columns=['SECTOR'])[etm_extract['YEAR'] == 2050].rename(columns={'VALUE': 'VALUE (TWh)'}).to_json(orient='index'), sub_task_id),
                                model=base_model
                            )
                        }
                    },
                    sector
                ): sector for sector in sectors
            }

            for future in concurrent.futures.as_completed(future_to_sector):
                sector = future_to_sector[future]
                try:
                    result = future.result()
                    if country not in etm_analysis:
                        etm_analysis[country] = {'country': country, 'tasks': {}}
                    etm_analysis[country]['tasks'].update(result)
                except Exception as e:
                    print(f'Error processing sector {sector}: {e}')
    else:
        for sector in sectors:
            sector_dict = etm_extract[etm_extract['SECTOR'] == sector]
            sector_dict = sector_dict.drop(columns=['SECTOR'])

            # Calculate percentages at sector level 
            sector_dict = sector_dict[sector_dict['YEAR'] == 2050]
            sector_dict.rename(columns={'VALUE': 'VALUE (TWh)'}, inplace=True)

            json_data_output = sector_dict.to_json(orient='index')
            context = data_analysis_report(instruction, json_data_output, sub_task_id)

            prompt = f"""Please give a concise summary of the data extracted. 
                        country: {country}. 
                        sector: {sector}.
                        data: {json_data_output}
                        Compare the data between 2019 and 2050.
                        """
            etm_results = oaicw(prompt, context, model=base_model)

            if country not in etm_analysis:
                etm_analysis[country] = {'country': country, 'tasks': {}}
            if sector not in etm_analysis[country]['tasks']:
                etm_analysis[country]['tasks'][sector] = {'sector': sector, 'final_copy': etm_results}

    return etm_analysis

def analyse_default_data(prompt_sheet, sub_task_id, detailed_context, input = 'default_charts'):
    row_extract = prompt_sheet['Default Charts'].set_index('sub_task_id', drop=False).loc[sub_task_id]
    description = row_extract['Description']
    chart_type = row_extract['Image']
    task_header = row_extract['Task Header']
    period_type = row_extract['Period_type']

    if chart_type == 'plexos_chart':
        plexos_extract = extract_plexos_data(row_extract, 'Default Charts', sub_task_id)
        plexos_extract = plexos_extract[plexos_extract['value'] != 0]
        for country in countries:
            try:
                plexos_results = process_plexos_data(country, plexos_extract, description, prompt_sheet, row_extract)
                detailed_context[country] = {f'Default Data - PLEXOS': plexos_results}
            except Exception as e:
                print(f'Error in analysing default data for {sub_task_id}', e, country)
                # error_log[sub_task_id] = {'function': 'Default Data', 'error': e, 'country': country}

    elif chart_type == 'Profile charts':
        collection = row_extract['collection_name']
        filename = fr'{base_location}\demand_profiles\{collection}.csv'
        data = pd.read_csv(filename)
        if period_type.lower() == 'monthly':
            data = data.groupby('Month').sum().reset_index()
            data = data[['Month', 'Value']]
            #normalise data
            data = data.div(data.sum(axis=1), axis=0)

        json_data_output = data.to_json(orient='index')
        instruction = f'Analyse the data. It show hourly demand profiles for {description}. The data is based on the data extracted from the {row_extract["DataBase"]} file'
        context = data_analysis_report(instruction, json_data_output, sub_task_id)
        prompt = f"Please give a concise summary of the data extracted from the profile charts"
        profile_results = oaicw(prompt, context, model=base_model)
        detailed_context = {f'Default Data - Hourly Profile': profile_results}

    elif chart_type == 'ETM':
        if run_concurrent:
            def process_country_etm(country):
                print('Analyzing default ETM data for', country)
                default_charts = row_extract
                etm_extract_country = extract_etm(sub_task_id, detailed_context, default_charts, country, description)
                context = main_context
                prompt = f"""Please give a concise summary of the data extracted from the ETM database.
                        Here are the Report extracts: {etm_extract_country}. return only the copy
                        """
                return oaicw(prompt, context, model=base_model)

            with concurrent.futures.ThreadPoolExecutor() as executor:
                futures = {executor.submit(process_country_etm, country): country for country in countries}
                for future in concurrent.futures.as_completed(futures):
                    country = futures[future]
                    try:
                        result = future.result()
                        detailed_context[country] = {f'Default Data - ETM': result}
                    except Exception as exc:
                        print(f'Error processing ETM for {country}:', exc)
        else:
            for country in countries:
                print('Analyzing default ETM data for', country)
                default_charts = row_extract
                etm_extract_country = extract_etm(sub_task_id, detailed_context, default_charts, country, description)
                context = main_context
                prompt = f"""Please give a concise summary of the data extracted from the ETM database.
                        Here are the Report extracts: {etm_extract_country}. return only the copy
                        """
                etm_result_analysis = oaicw(prompt, context, model=base_model)
                detailed_context[country] = {f'Default Data - ETM': etm_result_analysis}
    return detailed_context

def analyse_additional_data(prompt_sheet, sub_task_id, detailed_context, input='additional_data'):
    print('Analyzing additional data for', sub_task_id)
    guidelines = prompt_sheet['External Search'].loc[prompt_sheet['External Search'].index == sub_task_id]
    guidelines.set_index('Unique_ID', inplace=True)

    # Create separate dictionaries for country, level and title results
    country_results = {}
    level_results = {}
    title_results = {}

    all_additional_information_str = ' '.join(str(guidelines['additional information']))

    for guideline_id in guidelines.index:
        current_guideline = guidelines.loc[guideline_id]
        data_source = current_guideline['Data Source']
        level = current_guideline['Level'] 
        prompt = current_guideline['Prompt']
        description = current_guideline['Description']
        title = current_guideline['Title']

        if level == 'country':
            for country in countries:
                if country not in country_results:
                    country_results[country] = []

                if data_source == 'Internet Search':
                    print('Performing Internet Search for task:', sub_task_id, 'for country:', country)
                    additional_information = current_guideline['additional information']
                    additional_information = f'{additional_information} for {countries[country]}'
                    search_result = internet_search(country, additional_information, prompt, description)
                    country_results[country].append(search_result)

                elif data_source == 'ETM':
                    print(f'Extracting ETM data for {country}')
                    etm_results = extract_etm(sub_task_id, detailed_context, current_guideline, country, prompt)
                    country_results[country].append(etm_results)
                    
                elif data_source == 'PLEXOS':
                    plexos_extract = extract_plexos_data(current_guideline, 'External Search', sub_task_id)
                    plexos_results = process_plexos_data(country, plexos_extract, description, prompt_sheet, current_guideline)
                    country_results[country].append(plexos_results)

        else:
            if data_source == 'Internet Search':
                print('Analysing additional data for EU') 
                search_result = internet_search('EU', current_guideline['additional information'], prompt, description)
                level_results[level] = level_results.get(level, []) + [search_result]

            elif data_source == 'LLM Call':
                print('Analysing LLM data')
                llm_results = oaicw(title, description, model=base_model)
                title_results[title] = title_results.get(title, []) + [llm_results]

    # Synthesize results using LLM for each country
    for country in country_results:
        if country_results[country]:
            # Join all results into a single string and ensure it's a string type

            all_text = ' '.join(str(result) for result in country_results[country])
            synthesis_prompt = f"""
                                You are conducting research for the report and have extracted some data.
                                The title is {title}.
                                The description is {description}.
                                The prompt is {prompt}.
                                Here is the data that was researched: {all_additional_information_str}.
                                Here is the information that was obtained {level}: {all_text}
                                """
            country_synthesis = oaicw(synthesis_prompt, main_context, model=base_model)
            detailed_context[country] = {'synthesis': country_synthesis}

    # Synthesize level results 
    for level, results in level_results.items():
        if results:
            all_text = ' '.join(str(result) for result in results)
            synthesis_prompt = f"""
            Synthesize the following information for {level}: {all_text}
            """
            level_synthesis = oaicw(synthesis_prompt, main_context, model=base_model)
            detailed_context[country] = {'synthesis': level_synthesis}

    # Synthesize title results
    for title, results in title_results.items():
        if results:
            all_text = ' '.join(str(result) for result in results)
            synthesis_prompt = f"Synthesize the following information for {title}: {all_text}. Make it comprehensive and include the numbers"
            title_synthesis = oaicw(synthesis_prompt, main_context, model=base_model)
            detailed_context[country] = {'synthesis': title_synthesis}

    return detailed_context

def structure_dictionary(prompt_sheet, Report_structure, objective_id, task_id, sub_task_id, objective_title, task_title, sub_task_title, main_copy, summary, country):
    if objective_id not in Report_structure:
        Report_structure[objective_id] = {
            'Objective': objective_id,
            'tasks': {}
        }

    if task_id not in Report_structure[objective_id]['tasks']:
        Report_structure[objective_id]['tasks'][task_id] = {
            'task_id': task_id,
            'task_title': task_title,
            'sub_tasks': {}
        }

    if sub_task_id not in Report_structure[objective_id]['tasks'][task_id]['sub_tasks']:
        Report_structure[objective_id]['tasks'][task_id]['sub_tasks'][sub_task_id] = {
            'sub_task_id': sub_task_id,
            'sub_task_title': sub_task_title,
            'country': {}
        }

    if add_final_polish:
        final_copy =    refine_and_proofread(sub_task_id, main_copy, prompt_sheet, country, 'main')
    else:
        final_copy = main_copy

    if add_feedback and country != 'EU':
        feedback = feedback_sense_check(Report_structure[objective_id]['tasks'][task_id]['sub_tasks'][sub_task_id]['country'][country], main_context, main_copy)
        print('Feedback for', sub_task_id, 'for', country, feedback)
    else:
        feedback = None

    Report_structure[objective_id]['tasks'][task_id]['sub_tasks'][sub_task_id]['country'][country] = {
        'final_copy': final_copy,
        'objective_title': objective_title,
        'task_title': task_title,
        'sub_task_title': sub_task_title,
        'summary': None,
        'feedback': feedback
    }
    return Report_structure

def add_data_to_report(prompt_sheet, task_header, Report_structure, sub_task_id, objective_id, objective_title, task_title, sub_task_title, standard_guidelines, research_notes, detailed_context, additional_context, sub_task_description, location, task_id):
    country_context = f"""
                    {main_context}.
                    You are drafting section for objective: {objective_title}, task {task_title}, sub-task {sub_task_title}.
                    Consider you position in the document to determines how to write the section, espcially how the section starts and ends
                """
    country_prompt = f"""
                Write the report chapter {sub_task_id}: {task_header}, {sub_task_title} on the data analysis based on the information provided.
                Write a report on the data analysis for the country: {location}.
                Here is your instruction for the sub-task: {sub_task_description}.
                Here are the section guidelines: {standard_guidelines}.
                Here are the research notes: {research_notes}.
                Here is your context: {detailed_context}.
                Here is some data extracts to help you with the analysis: {additional_context}.   
                Illustrate with number if provided in the data.   
                You are writing a paragraph. Make it comprehensive, think about the connections in the data and how they are relevant to the chapter.
                Don't speak about policy or recommendations, just the data and the analysis, with some good insights on how the region is developing.
                Write in prose as it is a report, no additional formatting such as bold, italics or underlines.
                Do not include a header in the response, the text will be used directly directed to in the context.
                Only consider the country ranking if it is significant e.g. top 5 or bottom 5., else don't mention any ranking.
            """
    first_draft = oaicw(country_prompt, country_context, model = writer_model)
    first_draft = first_draft.replace('—', ', ')

    if add_summary:
        summary_context =   f"""
                                {main_context}.
                                You are asked to summarize the first draft of the report chapter for the europe.
                            """
        summary_prompt = f"""
                            Please summarize the first draft in a concise manner.
                            Here is the first draft: {first_draft}
                            """
        summary = oaicw(summary_prompt, summary_context, model = base_model)
    else:
        summary = None

    try:
        country = countries[location]
    except KeyError:
        country = location

    Report_structure = structure_dictionary(prompt_sheet, Report_structure, objective_id, task_id, sub_task_id, objective_title, task_title, sub_task_title, first_draft, summary, country)
    return Report_structure

def process_country_report(country, prompt_sheet, task_header, Report_structure, sub_task_id, objective_id, objective_title, task_title, sub_task_title, standard_guidelines, research_notes, detailed_context, additional_context, sub_task_description, task_id):
    """Helper function to process report for a single country"""
    print('Drafting section', sub_task_id, sub_task_title, 'for', country)
    try:
        detailed_context_country = detailed_context[country]
        try:
            additional_context_country = additional_context[country]
        except:
            additional_context_country = None
        
        return add_data_to_report(prompt_sheet, task_header, Report_structure, sub_task_id, objective_id, objective_title, task_title, sub_task_title,
                            standard_guidelines, research_notes, detailed_context_country, additional_context_country, sub_task_description, country, task_id)
    except Exception as e:
        print(f'Error in Drafting section {sub_task_id}', e, country)
        return None

def parse_sub_tasks(sub_tasks, prompt_sheet, objective_id, objective_title, task_title, Report_structure, task_id):
    for sub_task_id in sub_tasks.index:
        if not sub_task_list or sub_task_id in sub_task_list:
            sub_task_title = sub_tasks.loc[sub_task_id, 'Title']
            sub_task_description = sub_tasks.loc[sub_task_id, 'Description']
            geographical_level = sub_tasks.loc[sub_task_id, 'Geographic Level']
            task_header = sub_tasks.loc[sub_task_id, 'Task Header']
            print(sub_task_id, sub_task_description)

            detailed_context = {}
            additional_context = {}
                                
            standard_guidelines = prompt_sheet['Text Guidelines'].loc[sub_task_id, 'Standard']
            research_notes = prompt_sheet['Text Guidelines'].loc[sub_task_id, 'Research Notes']

            #extract information from the default charts
            for x in prompt_sheet['Default Charts']['sub_task_id']:
                if x == sub_task_id:
                    try:
                        detailed_context = analyse_default_data(prompt_sheet, sub_task_id, detailed_context, input='default_charts')
                    except Exception as e:
                        print(f'Error in analysing default data for {sub_task_id}', e)
                        detailed_context = None

            #extract information from the External Search
            if sub_task_id in prompt_sheet['External Search'].index.values:
                try:
                    additional_context = analyse_additional_data(prompt_sheet, sub_task_id, additional_context, input='additional_data')
                except Exception as e:
                    print(f'Error in analysing additional data for {sub_task_id}', e)
                    additional_context = None

            #draft the report for each country, the synthesis for the EU 
            if geographical_level == 'country':
                #find all countries in detailed_context where Default Data - PLEXOS is not None
                countries = [
                    country for country in detailed_context
                    if detailed_context[country].get('Default Data - PLEXOS') is not None
                    or detailed_context[country].get('Default Data - ETM') is not None
                ]
                if detailed_context is not None:
                    if run_concurrent:
                        # Process countries concurrently using ThreadPoolExecutor
                        with concurrent.futures.ThreadPoolExecutor() as executor:
                            future_to_country = {
                                executor.submit(
                                    process_country_report,
                                    country,
                                    prompt_sheet,
                                    task_header,
                                    Report_structure,
                                    sub_task_id,
                                    objective_id,
                                    objective_title,
                                    task_title,
                                    sub_task_title,
                                    standard_guidelines,
                                    research_notes,
                                    detailed_context,
                                    additional_context,
                                    sub_task_description,
                                    task_id
                                ): country for country in countries
                            }

                            for future in concurrent.futures.as_completed(future_to_country):
                                country = future_to_country[future]
                                try:
                                    result = future.result()
                                    if result:
                                        Report_structure.update(result)
                                except Exception as e:
                                    print(f'Error in Drafting section {sub_task_id}', e, country)
                    else:
                        # Process countries sequentially
                        for country in countries:
                            print('Drafting section', sub_task_id, sub_task_title, 'for', country)
                            try:
                                detailed_context_country = detailed_context[country]
                                try:
                                    additional_context_country = additional_context[country]
                                except:
                                    additional_context_country = None
                                
                                Report_structure = add_data_to_report(prompt_sheet, task_header, Report_structure, sub_task_id, objective_id, objective_title, task_title, sub_task_title, 
                                                    standard_guidelines, research_notes, detailed_context_country, additional_context_country, sub_task_description, country, task_id)
                            except Exception as e:
                                print(f'Error in Drafting section {sub_task_id}', e, country)

            else:
                print('Drafting section', sub_task_id, sub_task_title)
                Report_structure = add_data_to_report(prompt_sheet, task_header, Report_structure, sub_task_id, objective_id, objective_title, task_title, sub_task_title, 
                                    standard_guidelines, research_notes, detailed_context, additional_context, sub_task_description, geographical_level, task_id)      
        else:
            print('Skipping sub-task', sub_task_id)
    return Report_structure

def export_dictionary(Report_structure, objective_id):
    result_list = []
    for obj_id, objective_data in Report_structure.items():
        for task_id, task_data in objective_data['tasks'].items():
            for sub_task_id, sub_task_data in task_data['sub_tasks'].items():
                for country, country_data in sub_task_data['country'].items():
                    result_list.append({
                        'objective_id': obj_id,
                        'task_id': task_id,
                        'sub_task_id': sub_task_id,
                        'country': country,
                        'final_copy': country_data['final_copy'],
                        'objective_title': country_data['objective_title'],
                        'task_title': country_data['task_title'],
                        'sub_task_title': country_data['sub_task_title'],
                        'summary': country_data['summary']
                    })

    if result_list == []:
        result_df = pd.DataFrame(columns=['objective_id', 'task_id', 'sub_task_id', 'country', 'final_copy', 'objective_title', 'task_title', 'sub_task_title', 'summary'])
    else:
        result_df = pd.DataFrame(result_list)
        result_df = result_df[result_df['objective_id'] == objective_id]

    # Load existing results if available and merge missing sub_task_ids
    file_path = fr'{report_output_location}\{project_name}\objective_{objective_id}_results.csv'
    
    # Ensure the directory exists before saving
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    if os.path.exists(file_path):
        existing_df = pd.read_csv(file_path)
        missing_rows = existing_df[~existing_df['sub_task_id'].isin(result_df['sub_task_id'])]
        result_df = pd.concat([result_df, missing_rows], ignore_index=True)

    save_file(file_path, result_df, file_type = 'csv')    

def rebuild_report():
    Report_structure = pd.DataFrame()
    for objective in range(1,6):
        objective = pd.read_csv(fr'{report_output_location}\{project_name}\objective_{objective}_results.csv')
        Report_structure = pd.concat([Report_structure, objective])
    return Report_structure

def process_summary(index_row):
    index, row = index_row
    sub_task_id, country = index
    final_copy = row['final_copy']
    context =   f"""
                    {main_context}.
                    You are asked to summarize the first draft of the report chapter for the europe.
                """
    prompt = f"""
                Please summarize the first draft in a concise manner.
                write in prose with no headers or formatting
                Here is the first draft: {final_copy}
            """
    summary_copy = oaicw(prompt, context, model=base_model)
    summary_copy = summary_copy.replace('—', ', ')
    return summary_copy

def process_report_summary(country_copy, country):
    summary_prompt = f"""
                        If the country/region is "EU", this will be a default summary for the report.
                        Here is the data for summarisation: {country_copy}.
                        Please summarize the data in a concise manner, keeping it detailed enough to be a useful finalsummary.
                        The section will be used directly in the final report you return ONLY the final copy.
                        Keep the copy to around 400 words.
                    """
    summary_context = f"""
                        You are a copywriter drafting the final summary for the country/region {country}.
                        you must write in prose, with no headers or formatting. line breaks are allowed.
                        Note this section is the conclusion so tell a story of the developments in the country/region using simple language, in a professional way. 
                        Think the style of IEA world energy outlook reports.
                    """
    if 'gemini' in writer_model:
        try:
            summary_copy = gemini_call(summary_prompt, summary_context, model = writer_model)
        except Exception as e:
            print(f"Error in gemini_call: {e}")
            summary_copy = oaicw(summary_prompt, summary_context, model=sota_model)
    else:
        summary_copy = oaicw(summary_prompt, summary_context, model=sota_model)
    return summary_copy

def process_sub_task_summary(sub_task_id, report_section):
    sub_task_summary = report_section[report_section['sub_task_id'] == sub_task_id]
    objective_id = sub_task_summary['objective_id'].iloc[0]
    task_id = sub_task_summary['task_id'].iloc[0]
    sub_task_id = sub_task_summary['sub_task_id'].iloc[0]
    objective_title = sub_task_summary['objective_title'].iloc[0]
    task_title = sub_task_summary['task_title'].iloc[0]
    sub_task_title = sub_task_summary['sub_task_title'].iloc[0]

    sub_task_summary = sub_task_summary.drop(columns=['objective_id', 'task_id', 'final_copy', 'objective_title', 'task_title', 'sub_task_title'])
    sub_task_summary.reset_index(drop=True, inplace=True)
    sub_task_summary_dict = sub_task_summary.to_dict(orient='index')
    summary_prompt = f"""
                        You are drafting introduction section for a section of the report.                    
                        Please create a comprehensive summary of this data {sub_task_summary_dict}
                        Please include any key figures, comparison across countries in the data.
                    """
    summary_context = f"""
                        {main_context}.
                        write in prose, with no headers or formatting.
                        Note this section is the conclusion so tell a story of the development in that country using simple language.
                        Return only the final copy as it will be used directly in the report.
                    """

    if 'gemini' in writer_model:
        summary_copy = gemini_call(summary_prompt, summary_context, model = writer_model)
    else:
        summary_copy = oaicw(summary_prompt, summary_context, model=sota_model)
    
    return [sub_task_id, 'EU', objective_id, task_id, summary_copy, objective_title, task_title, sub_task_title, summary_copy]

def process_country_summary(country):
    print('Summarizing the report for', country)
    country_name = countries[country]
    country_summary = updated_structure[updated_structure['country'] == country_name]
    country_summary.reset_index(drop=True, inplace=True)
    country_summary_dict = country_summary.to_dict(orient='index')
    summary_prompt = "Please summarize the data in a concise manner, keeping it detailed enough to be a useful summary."
    summary_context = f"""
                        You are drafting the report for the country: {country_name}. 
                        Here is the data from the sections of the report: {country_summary_dict}
                        write in prose.
                        Note this section is the conclusion so tell a story of the development in that country using simple language.
                        """
    if 'gemini' in writer_model:
        gemini_call(summary_prompt, summary_context, model = writer_model)
    else:
        summary_copy = oaicw(summary_prompt, summary_context, model=sota_model)
    return country_name, summary_copy

def conclusion_next_steps(prompt_sheet):
    print('Summarizing the report')
    #for each row in the analysis dataframe, summarize the copy in the 'final_copy' column using a oaicw
    for objective in summary_objective_list:
        report_section = pd.read_csv(fr'{report_output_location}\{project_name}\objective_{objective}_results.csv')
        sub_task_ids = report_section['sub_task_id'].unique()

        #this section will add a summary to every sub_task copy
        if add_sub_task_summary:
            report_section.set_index(['sub_task_id', 'country'], inplace=True)
            with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                future_to_index = {executor.submit(process_summary, item): item[0] for item in report_section.iterrows()}
                for future in concurrent.futures.as_completed(future_to_index):
                    try:
                        summary_copy = future.result()
                        report_section.loc[(sub_task_id, country), 'summary'] = summary_copy
                    except Exception as e:
                        print(f"Error processing summary: {e}")
            report_section.reset_index(inplace=True)

        #this part will create the EU summary
        if add_introduction_to_task:
            for sub_task_id in sub_task_ids:
                try:
                    result = process_sub_task_summary(sub_task_id, report_section)
                    report_section.loc[len(report_section)] = result
                except Exception as e:
                    print(f"Error processing sub-task {sub_task_id}: {e}")
            report_section.to_csv(fr'{report_output_location}\{project_name}\objective_{objective}_results.csv', index=False)

    #This section will write section the summary
    if draft_report_conclusion:
        report_section = pd.DataFrame()
        for objective in summary_objective_list:
            report_section = pd.concat([report_section, pd.read_csv(fr'{report_output_location}\{project_name}\objective_{objective}_results.csv')])   
        countries = report_section['country'].unique()
        report_section.set_index(['country'], inplace=True)

        summary_section = pd.DataFrame(columns=['sub_task_id', 'country', 'objective_id', 'task_id', 'final_copy', 'objective_title', 'task_title', 'sub_task_title', 'summary'])

        for country in countries:
            country_text_all = report_section.loc[country]
            country_summaries = country_text_all[['final_copy', 'objective_title', 'task_title', 'sub_task_title']]
            country_summaries.reset_index(drop=True, inplace=True)
            country_summaries_dict = country_summaries.to_dict(orient='index')
            summary_copy = process_report_summary(country_summaries_dict, country)
            summary_section.loc[len(summary_section) + 1] = ['5.1.1', country, 5, 5.1, summary_copy, 'Conclusion', 'Conclusion', 'Conclusion', summary_copy]
            print('Report summarised for', country)
        summary_section.to_csv(fr'{report_output_location}\{project_name}\objective_5_results.csv', index=False)

    report_section = pd.read_csv(fr'{report_output_location}\{project_name}\objective_5_results.csv')
    print('Drafting the next steps for the report')
    standard_guidelines = prompt_sheet['Text Guidelines'].loc['5.2.1', 'Standard']
    next_steps_prompt = f"""
                            Please draft the section on next steps. Please write in prose.
                            Here are the notes for the next steps: {standard_guidelines}.
                            Return the final copy ONLY as it will be used directly in the report.
                            Do not add any headers or formatting.
                        """
    next_steps_context = f"""
                            {main_context}.
                        """

    if 'gemini' in writer_model:
        next_steps_copy = gemini_call(next_steps_prompt, next_steps_context, model = writer_model)
    else:
        next_steps_copy = oaicw(next_steps_prompt, next_steps_context, model=sota_model)

    report_section.loc[len(report_section) + 1] = ['5.2.1', 'EU', 5, 5.2, next_steps_copy,  'Next Steps', 'Next Steps', 'Next Steps', next_steps_copy]
    report_section.to_csv(fr'{report_output_location}\{project_name}\objective_5_results.csv', index=False)
    print('Report summarization complete')
    # return Summary_Report_structure

def parse_objectives(prompt_sheet, objective_id, Report_structure, context):
    objective_title = prompt_sheet['Objectives'].loc[objective_id, 'Title']
    objective_description = prompt_sheet['Objectives'].loc[objective_id, 'Description']
    print(objective_id, objective_description)
    context['Objective'] = objective_description
    tasks = prompt_sheet['Tasks'][prompt_sheet['Tasks']['Objective ID'] == objective_id]

    for task_id in tasks.index:
        if task_id > 0:   
            task_title = prompt_sheet['Tasks'].loc[task_id, 'Title']
            task_description = prompt_sheet['Tasks'].loc[task_id, 'Description']
            context['Task'] = task_description
            print(task_id, task_description)
            sub_tasks = prompt_sheet['Sub Tasks'][prompt_sheet['Sub Tasks']['Task_section_id'] == task_id]
            Report_structure = parse_sub_tasks(sub_tasks, prompt_sheet, objective_id, objective_title, task_title, Report_structure, task_id)
    
    export_dictionary(Report_structure, objective_id)
    return Report_structure

def check_report_objectives(objectives):
    report_structure = {}
    for objective_id in objectives:
        if objective_id not in report_structure:
            file_path = fr'{report_output_location}\{project_name}\objective_{objective_id}_results.csv'
            if os.path.exists(file_path):
                result_df = pd.read_csv(file_path)
                result_dict = result_df.to_dict(orient='index')
                report_structure.update(result_dict)
            else:
                print(f'Warning: Objective {objective_id} results file not found at {file_path}')
    return report_structure

def main(prompt_sheet):
    global objective_list
    Report_structure  = {}
    for aim_id in prompt_sheet['Aims'].index:
        context = {}
        aim_description = prompt_sheet['Aims'].loc[aim_id, 'Description']
        context['Aim'] = aim_description
        print(aim_id, aim_description)
        objectives = prompt_sheet['Objectives'].index.unique()
        if objective_list == []:
            objective_list = objectives.tolist()

        if run_main_report:
            if run_concurrent == True:
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future_objectives = {executor.submit(parse_objectives, prompt_sheet, objective_id, Report_structure, context): objective_id for objective_id in objective_list}
                    for future in concurrent.futures.as_completed(future_objectives):
                        objective_id = future_objectives[future]
                        result = future.result()
                        Report_structure.update(result)
            else:
                for objective_id in objectives:
                    if objective_id in objective_list:
                        print('Parsing objective', objective_id)
                        Report_structure = parse_objectives(prompt_sheet, objective_id, Report_structure, context)

            check_report_objectives(objectives)

        if conclude_report:
            conclusion_next_steps(prompt_sheet)

    final_report = rebuild_report()
    final_report.drop(columns = ['summary'], inplace = True)

    file_path = fr'{report_output_location}\{project_name}\{project_name}.csv'
    save_file(file_path, final_report, file_type = 'csv')
    create_word_file(final_report, output_path=fr'{report_output_location}\Joule_Report_V2.docx')

def run_prompt_sheet():
    start_timer = time.time()
    start_time = time.strftime('%H:%M:%S', time.localtime())
    print(start_time)
    warnings.filterwarnings('ignore')
    prompt_sheet = eps.import_excel_sheets_as_dict(file_path)
    main(prompt_sheet)
    end_timer = time.time()
    end_time = time.strftime('%H:%M:%S', time.localtime())
    print('Time taken:', (end_timer - start_timer)/3600,'hours')
    print('Start time:', start_time, 'End time:', end_time)

project = 'Joule'

if project == 'Joule':
    file_path = r'C:\Users\ENTSOE\Tera-joule\Terajoule - Terajoule\Projects\Sectoral Model\Website\Joule_prompt_sheet_.xlsx'

if project == 'Medtso':
    file_path = r'C:\Users\ENTSOE\Tera-joule\Terajoule - Terajoule\Projects\Sectoral Model\Website\TEASIMED_IoSN_prompt_sheet.xlsx'

etm = pd.read_csv(r'external_resources\model_databases\joule_model\etm_2024\etm_extract_joule_report.csv')
nodalsplit = pd.read_excel(r'src\EMIL\demand\Input\Population and Industrial Sizes.xlsx', sheet_name = 'ehighway_nut2 conversion')
base_location = r'external_resources\model_databases\joule_model'
report_output_location = fr'functions\copywriting_functions\Generated Reports\{project}'
#create folder if it doesn't exist
os.makedirs(fr'{report_output_location}', exist_ok=True)

prompt_sheet = eps.import_excel_sheets_as_dict(file_path)

tone_and_feel = {
                'purpose': 'Educate',
                'target_audience': 'Policy Makers, TSO, Energy Project Owner, NGOs',
                'tone_voice': 'Slightly pro nuclear',
                'format': 'Energy Report',
                }
countries = {
            "AT": "Austria", 
            "BE": "Belgium", "BG": "Bulgaria", "CY": "Cyprus", "CZ": "Czech Republic", 
            "DE": "Germany", "DK": "Denmark", "EE": "Estonia", "ES": "Spain", "FI": "Finland", "FR": "France", "GR": "Greece", 
            "HR": "Croatia", "HU": "Hungary", "IE": "Ireland", "IT": "Italy", "LT": "Lithuania", "LU": "Luxembourg", "LV": "Latvia", 
            "MT": "Malta", "NL": "Netherlands", "PL": "Poland", "PT": "Portugal", "RO": "Romania",  "SE": "Sweden", "SI": "Slovenia", 
            "SK": "Slovakia", "UK": "United Kingdom"
            }

test_mode = False
local_mode = False
open_source_mode = True
open_source_sota_mode = False
closed_source_mode = False

run_concurrent = False
skip_internet_search = False
add_summary = False
add_feedback = False
run_main_report = True
conclude_report = True
add_final_polish = False
add_sub_task_summary = False
add_introduction_to_task = False 
draft_report_conclusion = False
search_model = 'sonar'

model_name = 'TJ Dispatch_Future_Nuclear+'
references = {}
objective_list = []
sub_task_list = []
summary_objective_list = []

country_borders_file = r'config\topology\country_borders.yaml'
with open(country_borders_file, 'r') as f:
    country_borders = yaml.safe_load(f) or {}

main_context = str(prompt_sheet['Read me']['Context'])
project_name = str(prompt_sheet['Read me']['Project Name'][0])

nodalsplit['Country'] = nodalsplit['Node'].str[0:2]
nodalsplit.drop(columns = ['Position'], inplace = True)

default_ai_models_file = r'config\default_ai_models.yaml'
with open(default_ai_models_file, 'r') as f:
    ai_models_config = yaml.safe_load(f)
base_model = ai_models_config.get("base_model", "gpt-5-mini")
pro_model = ai_models_config.get("pro_model", "gpt-5")
sota_model = ai_models_config.get("sota_model", "gpt-5")
writer_model = ai_models_config.get("writer_model", "gpt-5")

if __name__ == "__main__": 
    run_prompt_sheet()
