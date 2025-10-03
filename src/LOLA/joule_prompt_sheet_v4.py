# -*- coding: utf-8 -*-
"""
Joule Prompt Sheet Executor V4 - Object-Oriented Implementation
Created: October 3, 2025
@author: AI Architecture Team
"""

import sys
import os
import pandas as pd
import warnings
import re
import time
import json
import concurrent.futures
from threading import Lock
from fpdf import FPDF
from docx import Document
import yaml
from typing import Dict, List, Optional, Any, Tuple
from dataclasses import dataclass, field
from abc import ABC, abstractmethod

# Path configuration
top_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), os.pardir, os.pardir))
if top_dir not in sys.path:
    sys.path.insert(0, top_dir)

sys.path.append('utils')
sys.path.append('functions/plexos_functions')
sys.path.append('functions')
sys.path.append('functions/LLMs')

import src.LOLA.extract_excel_sheets as eps
from src.ai.llm_calls.open_ai_calls import run_open_ai_ns as roains
from src.ai.llm_calls.open_ai_calls import open_ai_search
from src.ai.llm_calls.gemini import google_search


def llm_call_with_json(prompt: str, context: str, model: str, extract_key: str = 'copy') -> str:
    """
    Wrapper for LLM calls that handles JSON response format.
    
    Args:
        prompt: The user prompt
        context: The system context
        model: The model to use
        extract_key: The JSON key to extract from response (default: 'copy')
    
    Returns:
        Extracted text from the JSON response
    """
    json_prompt = f"""{prompt}
                            
                            IMPORTANT: Return your response as a JSON object with the following structure:
                            {{
                                "copy": "your main response text here",
                                "reasoning": "brief explanation of your analysis and approach"
                            }}
                    """
                            
    try:
        response = roains(json_prompt, context, model=model)
        
        # Try to parse as JSON
        import json
        try:
            response_json = json.loads(response)
            return response_json.get(extract_key, response_json.get('copy', response))
        except json.JSONDecodeError:
            # If JSON parsing fails, return the raw response
            print(f"Warning: Could not parse JSON response, returning raw text")
            return response
    except Exception as e:
        print(f"Error in llm_call_with_json: {e}")
        raise


@dataclass
class ProjectConfig:
    """Configuration settings for the report generation project."""
    project_name: str
    file_path: str
    base_location: str
    report_output_location: str
    model_name: str
    tone_and_feel: Dict[str, str]
    countries: Dict[str, str]
    main_context: str
    
    # Model configurations
    base_model: str = "gpt-5-mini"
    pro_model: str = "gpt-5"
    sota_model: str = "gpt-5"
    writer_model: str = "gpt-5"
    search_model: str = "google"
    
    # Feature flags
    test_mode: bool = False
    local_mode: bool = False
    open_source_mode: bool = True
    run_objectives_concurrent: bool = True
    run_task_concurrent: bool = True
    skip_internet_search: bool = False
    add_summary: bool = False
    add_feedback: bool = False
    run_main_report: bool = True
    conclude_report: bool = True
    add_final_polish: bool = False
    add_sub_task_summary: bool = False
    add_introduction_to_task: bool = False
    draft_report_conclusion: bool = False
    
    # Processing lists
    objective_list: List[int] = field(default_factory=list)
    sub_task_list: List[str] = field(default_factory=list)
    summary_objective_list: List[int] = field(default_factory=list)
    
    country_borders: Dict[str, List[str]] = field(default_factory=dict)


class FileManager:
    """Handles file operations including saving and reading."""
    
    @staticmethod
    def save_file(file_path: str, doc: Any, file_type: str = "word") -> None:
        """Save a file with retry logic for permission errors."""
        if file_type == "word":
            while True:
                try:
                    doc.save(file_path)
                    break
                except PermissionError:
                    print("Please close the file to continue")
                    file_close = input("Is the file closed? (yes/no): ")
                    if file_close.lower() == "yes":
                        break
                    time.sleep(5)

        elif file_type == 'csv':
            while True:
                try:
                    doc.to_csv(file_path, index=False)
                    break
                except PermissionError:
                    print("Please close the file to continue")
                    file_close = input("Is the file closed? (yes/no): ")
                    if file_close.lower() == "yes":
                        break
                    time.sleep(5)
    
    @staticmethod
    def ensure_directory_exists(directory_path: str) -> None:
        """Create directory if it doesn't exist."""
        os.makedirs(directory_path, exist_ok=True)


class ReportExporter:
    """Handles export of reports to various formats."""
    
    def __init__(self, file_manager: FileManager):
        self.file_manager = file_manager
    
    def create_pdf(self, final_report: pd.DataFrame, output_path: str = "Joule_Report.pdf") -> None:
        """Create PDF from final report."""
        print("Creating PDF from final report...")
        pdf = FPDF()
        pdf.set_font("Arial", size=12)
        pdf.add_page()
        
        for _, row in final_report.iterrows():
            pdf.set_font("Arial", "B", 16, uni=True)  # H1
            pdf.multi_cell(0, 10, str(row.get("objective_id", "")))

            pdf.set_font("Arial", "B", 14, uni=True)  # H2
            pdf.multi_cell(0, 10, str(row.get("task_id", "")))

            pdf.set_font("Arial", "B", 12, uni=True)  # H3
            pdf.multi_cell(0, 10, str(row.get("sub_task_id", "")))

            pdf.set_font("Arial", "B", 10, uni=True)  # H4
            pdf.multi_cell(0, 10, str(row.get("country", "")))

            pdf.set_font("Arial", "", 11, uni=True)   # Paragraph
            pdf.multi_cell(0, 10, str(row.get("final_copy", "")))

        try:
            pdf.output(output_path)
            print("PDF created successfully!")
            os.startfile(output_path)
        except Exception as e:
            print("Error creating PDF:", e)
    
    def create_word_file(self, final_report: pd.DataFrame, output_path: str) -> None:
        """Create Word document from final report."""
        print("Creating Word file from final report...")
        doc = Document()

        for obj_id, obj_group in final_report.groupby("objective_id"):
            obj_title = obj_group["objective_title"].iloc[0]
            doc.add_heading(f"{obj_id} {obj_title}", level=1)

            for t_id, task_group in obj_group.groupby("task_id"):
                t_title = task_group["task_title"].iloc[0]
                if t_title != obj_title:
                    doc.add_heading(f"{t_id} {t_title}", level=2)

                for s_id, sub_task_group in task_group.groupby("sub_task_id"):
                    s_title = sub_task_group["sub_task_title"].iloc[0]
                    if s_title != t_title:
                        doc.add_heading(f"{s_id} {s_title}", level=3)

                    for country, country_group in sub_task_group.groupby("country"):
                        if country != "EU" and country != "summary" and country != s_title:
                            doc.add_heading(country, level=4)

                        for _, row in country_group.iterrows():
                            final_copy = str(row.get("final_copy", ""))
                            doc.add_paragraph(final_copy)

        self.file_manager.save_file(output_path, doc, file_type="word")
        print("Word file created successfully!")


class DataExtractor(ABC):
    """Abstract base class for data extraction operations."""
    
    def __init__(self, config: ProjectConfig):
        self.config = config
    
    @abstractmethod
    def extract(self, *args, **kwargs) -> Any:
        """Extract data from source."""
        pass


class PlexosDataExtractor(DataExtractor):
    """Extracts and processes PLEXOS simulation data."""
    
    def extract_csv_data(self, collection: str, interval_name: str, 
                        category: Optional[str] = None, 
                        property: Optional[str] = None) -> pd.DataFrame:
        """Extract CSV data from PLEXOS output."""
        if interval_name.lower() == 'yearly':
            location = os.path.join(self.config.base_location, interval_name.lower(), 
                                   self.config.model_name, f'{collection.lower()}.csv')
            data_extract = pd.read_csv(location)

            if category != '-':
                try:
                    valid_categories = [cat.strip() for cat in category.split(',')]
                    data_extract = data_extract[data_extract['category_name'].isin(valid_categories)]
                except:
                    if category is not None:
                        data_extract = data_extract[data_extract['category_name'] == category]
                    if property is not None:
                        data_extract = data_extract[data_extract['property_name'] == property]

            data_extract = data_extract[['collection_name', 'category_name', 'child_name', 
                                        'property_name', 'value', 'unit_name', 'interval_id']]

        elif interval_name == 'daily':
            data_extract = pd.DataFrame()
            try:
                categories = [cat.strip() for cat in category.split(',')]
                for cat in categories:
                    location = os.path.join(self.config.base_location, interval_name.lower(), 
                                          self.config.model_name,
                                          f'{collection.lower()}_{cat.lower()}_{property.lower()}.csv')
                    all_data = pd.read_csv(location)
                    data_extract = pd.concat([data_extract, all_data], ignore_index=True)
                    
                    # Convert to monthly data
                    monthly_data_extract = data_extract.copy()
                    monthly_data_extract['month_year'] = pd.to_datetime(
                        monthly_data_extract['date_string'], 
                        dayfirst=True, 
                        format="%d/%m/%Y"
                    ).dt.to_period('M')
                    
                    group_cols = [col for col in monthly_data_extract.columns 
                                if col not in ['value', 'date_string', 'interval_id', 'period_value']]
                    monthly_data_extract = monthly_data_extract.groupby(group_cols).sum().reset_index()
                    monthly_data_extract.drop(columns=['interval_id', 'period_value', 'date_string'], 
                                             inplace=True)
                    data_extract = monthly_data_extract
            except Exception as e:
                location = os.path.join(self.config.base_location, interval_name.lower(), 
                                       self.config.model_name,
                                       f'{collection.lower()}_{category.lower()}_{property.lower()}.csv')
                data_extract = pd.read_csv(location)
        
        return data_extract
    
    def extract(self, prompt_sheet: Dict, sheet: str, sub_task_id: str) -> pd.DataFrame:
        """Extract PLEXOS data based on prompt sheet configuration."""
        collection_name = prompt_sheet['collection_name']
        property_list = prompt_sheet['property_list']
        category_list = prompt_sheet['category_list']
        interval_name = prompt_sheet['Period_type']

        plexos_extract = self.extract_csv_data(collection_name, interval_name, 
                                               category=category_list, property=property_list)
        plexos_extract = plexos_extract[plexos_extract['value'] != 0]
        plexos_extract['country'] = plexos_extract['child_name'].str[0:2]
        return plexos_extract


class ETMDataExtractor(DataExtractor):
    """Extracts and processes ETM (Energy Transition Model) data."""
    
    def __init__(self, config: ProjectConfig, etm_data: pd.DataFrame):
        super().__init__(config)
        self.etm_data = etm_data
    
    def extract(self, sub_task_id: str, detailed_context: Dict, df: pd.DataFrame, 
                country: str, instruction: str) -> Dict:
        """Extract and analyze ETM data for a specific country."""
        df.fillna('-', inplace=True)
        etm_extract = self.etm_data[self.etm_data['VALUE'] != 0]
        etm_extract = etm_extract.drop(columns=['PROPERTY', 'UNIT'])
        etm_analysis = {}

        etm_extract = etm_extract[etm_extract['COUNTRY'] == country]
        sectors = etm_extract['SECTOR'].unique()

        for sector in sectors:
            sector_dict = etm_extract[etm_extract['SECTOR'] == sector]
            sector_dict = sector_dict.drop(columns=['SECTOR'])
            sector_dict = sector_dict[sector_dict['YEAR'] == 2050]
            sector_dict.rename(columns={'VALUE': 'VALUE (TWh)'}, inplace=True)

            json_data_output = sector_dict.to_json(orient='index')
            context = self._build_analysis_context(instruction, json_data_output, sub_task_id)

            prompt = f"""Please give a concise summary of the data extracted. 
                        country: {country}. 
                        sector: {sector}.
                        data: {json_data_output}
                        Compare the data between 2019 and 2050. 
                        """
            etm_results = llm_call_with_json(prompt, context, model=self.config.base_model)

            if country not in etm_analysis:
                etm_analysis[country] = {'country': country, 'tasks': {}}
            if sector not in etm_analysis[country]['tasks']:
                etm_analysis[country]['tasks'][sector] = {
                    'sector': sector, 
                    'final_copy': etm_results
                }

        return etm_analysis
    
    def _build_analysis_context(self, instruction: str, json_data: str, chapter: str) -> str:
        """Build context for data analysis."""
        return f"""{self.config.main_context}. 
                You are currently drafting chapter {chapter}.
                Instruction: {instruction}.
                Do not end the chapter with any conclusion, unless it is stated we are at the conclusions chapter.
                Do not include a header in the response, the text will be used directly directed to in the context.
                Do neutral language, do not overuse adjectives. Use simple language.
                """


class InternetSearchService:
    """Handles internet search operations for data gathering."""
    
    def __init__(self, config: ProjectConfig):
        self.config = config
    
    def search(self, country: str, kpi: str, title: str, description: str) -> str:
        """Perform internet search for specific data."""
        context = f"""
                    {self.config.main_context}.
                    You are asked to analyse some additional data.
                    you must perform an internet search to obtain the data. focus on the year 2024.
                    The KPI is {kpi}. The title is {title}. The description is {description}.
                    """
        prompt = f""" 
                    Please perform an internet search to obtain the data.
                    The KPI is {kpi}. The title is {title}. The description is {description}. 
                    Focus on the country {country}. 
                    """
        

        search_results = google_search(prompt)

        return search_results


class TextRefinementService:
    """Handles text refinement and proofreading operations."""
    
    def __init__(self, config: ProjectConfig):
        self.config = config
    
    def refine_and_proofread(self, sub_task_id: str, current_text: str, 
                            prompt_sheet: Dict, country: str, 
                            section_type: str = 'main') -> str:
        """Refines and proofreads text for report quality."""
        try:
            reference_text = prompt_sheet['Cross Reference'].loc[sub_task_id, 'Copy']
        except KeyError:
            reference_text = None

        refinement_prompt = f"""
                                    Follow these guidelines strictly:

                                    1. Maintain a consistent tone aligned with: {self.config.tone_and_feel}.
                                    2. Use varied sentence structures; avoid fluff, marketing speak, or repetitive phrasing.
                                    3. Proofread thoroughly for grammar, spelling, and clarity.
                                    4. Write in prose suitable for a report; do not add formatting such as bold, italics, or underlines.
                                    5. Since this is not an introductory section, avoid phrases like "this report," "this section," or "this subsection."
                                    6. Craft the text to mimic human writing that would pass an AI detection test.

                                    Authorship guidelines to adhere to:
                                    - Statistical predictability (perplexity): Ensure the text includes spikes of randomness and surprise, avoiding overly smooth or predictable AI-like patterns.
                                    - Uniformity of structure: Vary sentence and paragraph lengths; avoid consistent rhythm or pacing.
                                    - Lexical fingerprints: Do not overuse AI-favored tokens or stock phrases (e.g., "delve," "landscape," "pivotal," "robust framework"); avoid repetition across contexts.
                                    - Probability curvature anomalies: Write in a way that deviates from typical AI probability basins.
                                    - Watermarking signals: If applicable, avoid hidden token patterns associated with AI models.

                                    Current text to refine:
                                    {current_text}

                                    Reference style example:
                                    {reference_text}

                                    Country: {country}
                                    """

        context = f"""
                    You are a professional editor writing in the style of IRENA and the International Energy Agency. 
                    You are a historian writing a report on the energy landscape in 2050.
                    Never include the chapter name or heading in the response, this will destroy the flow of the whole report.
                    Write in prose, do not use bullet points or lists. Do not use any formatting such as bold, italics or underlines.
                    Do not create paragraph breaks, keep the text as one or 2 continuous paragraphs.
                    You must ensure the text is engaging, whilst remaining professional and informative, as it is a technical report.
                    Avoid perfect grammar, use simple language, avoid jargon, and keep the language simple and natural.
                    Add some contractions. Avoid repetitive phrases. 
                    Avoid Generic Openings such as "In today's world...", "It is important to...", "In Chapter x we discuss".
                    Keep language simple, natural and authoritative.
                    Finally return ONLY the final copy as it will be used verbatim in the report.
        """

        try:
            refined_text = llm_call_with_json(refinement_prompt, context, model=self.config.writer_model)
        except Exception as e:
            print(f"Error in refining text for {sub_task_id}: {e}. Running with o3-mini")
            refined_text = llm_call_with_json(refinement_prompt, context, model='o3-mini')

        refined_text = refined_text.replace('—', ', ')
        return refined_text
    
    def feedback_sense_check(self, report_extract_dict: Dict, context: str, main_copy: str) -> str:
        """Perform sense check on report section."""
        objective_title = report_extract_dict['objective_title']
        task_title = report_extract_dict['task_title']
        sub_task_title = report_extract_dict['sub_task_title']
        country = report_extract_dict['country']

        sense_check_context = f"""
                    {context}.
                    You are asked to sense check the following section of the report. 
                    The year is 2050 so the energy landscape will have changed compared to today.
                    """

        prompt = f"""
                {sense_check_context}
                You are asked to sense check the following section of the report:
                Objective: {objective_title}
                Task: {task_title}
                Sub-task: {sub_task_title}
                Country: {country}
                Information: {main_copy}
                """
        feedback = llm_call_with_json(prompt, sense_check_context, model=self.config.sota_model)
        return feedback


class DataAnalysisService:
    """Handles data analysis and processing operations."""
    
    def __init__(self, config: ProjectConfig, plexos_extractor: PlexosDataExtractor,
                 nodalsplit: pd.DataFrame):
        self.config = config
        self.plexos_extractor = plexos_extractor
        self.nodalsplit = nodalsplit
    
    def get_ranking(self, df: pd.DataFrame, country: str, 
                   value_col: str, country_col: str) -> str:
        """Get country ranking based on data."""
        df = df.groupby([country_col]).sum().reset_index()
        df['rank'] = df[value_col].rank(ascending=False)
        country_rank = df[df[country_col] == country]['rank'].values[0]
        country_rank_text = (f'{self.config.countries[country]} is ranked {int(country_rank)} '
                           f'out of {len(df)} countries')
        return country_rank_text
    
    def process_plexos_data(self, country: str, plexos_extract: pd.DataFrame,
                           description: str, prompt_sheet: Dict, 
                           row_extract: Dict) -> Optional[str]:
        """Process PLEXOS data for a specific country."""
        print('Analyzing PLEXOS data for', country)
        property_name = row_extract['property_list']
        categories = row_extract['category_list']
        sub_task_header = row_extract['sub_task_header']

        try:
            additional_information = row_extract['additional information']
        except:
            additional_information = None

        plexos_extract_df = plexos_extract[plexos_extract['property_name'] == property_name]
        
        if categories != '-':
            try:
                valid_categories = [cat.strip() for cat in categories.split(',')]
                plexos_extract_df = plexos_extract_df[
                    plexos_extract_df['category_name'].isin(valid_categories)
                ]
            except:
                if categories is not None:
                    plexos_extract_df = plexos_extract_df[
                        plexos_extract_df['category_name'] == categories
                    ]

        category_df = prompt_sheet['All Categories'][['category_name', 'group', 'Description']]
        plexos_extract_df = plexos_extract_df.merge(category_df, on='category_name', how='left')

        if 'Grid Flows' in sub_task_header:
            bordering_countries = self.config.country_borders.get(country, [])
            bordering_countries.append(country)
            plexos_extract_country = plexos_extract_df[
                plexos_extract_df['country'].isin(bordering_countries)
            ]
        else:
            plexos_extract_country = plexos_extract_df[plexos_extract_df['country'] == country]
        
        plexos_extract_country['region'] = plexos_extract_country['child_name'].str[2:4]
        plexos_extract_country['region'] = plexos_extract_country['region'].apply(
            lambda x: 'Entire Country' if x == '00' else x.lstrip('0')
        )

        units = plexos_extract_country['unit_name'].unique()
        asset_classes = plexos_extract_country['collection_name'].unique()

        try:
            plexos_extract_country.drop(
                columns=['category_name', 'Description', 'interval_id', 
                        'country', 'property_name', 'unit_name'], 
                inplace=True
            )
        except:
            plexos_extract_country.drop(
                columns=['model_name', 'collection_name', 'category_name', 'Description',
                        'country', 'property_name', 'unit_name'], 
                inplace=True
            )

        plexos_extract_country.rename(columns={'group': 'category'}, inplace=True)
        plexos_extract_country = plexos_extract_country.drop_duplicates(keep='first')

        try:
            plexos_extract_country = plexos_extract_country.groupby(
                ['child_name', 'category', 'region']
            ).sum().reset_index()
        except:
            pass

        try:
            json_data_output = plexos_extract_country.to_json(orient='index')
        except:
            json_data_output = plexos_extract_country.to_dict(orient='records')

        try:
            country_rank = self.get_ranking(plexos_extract_df, country, 
                                          value_col='value', country_col='country')
        except:
            country_rank = None

        city_names = self.nodalsplit[self.nodalsplit['Country'] == country]
        city_names['region'] = city_names['Node'].str[2:4]
        city_names['region'] = city_names['region'].apply(
            lambda x: 'Entire Country' if x == '00' else x.lstrip('0')
        )
        city_names = city_names.drop(columns=['Node'])
        city_names_json = city_names.to_json(orient='index')

        default_data_context = self._build_data_analysis_context(description, 
                                                                 json_data_output, 
                                                                 description)
        
        default_data_prompt = f"""Please give a concise summary of the data extracted. 
                                Give a neutral tone and use simple non emotional language.
                                {additional_information}.
                                Here are the city names: {city_names_json}. 
                                Consider the geographical location e.g. north, east, south west etc.
                                Start by giving a general overview of the regions, their cities and location.
                                Here are the results: {json_data_output}.
                                Units: {units}.
                                Country: {country}. 
                                Properties: {property_name}.
                                Categories: {categories}.
                                Asset classes: {asset_classes}.
                                Country ranking: {country_rank}.
                                Do not perform any formatting e.g. bold, italics.
                                """
        
        if plexos_extract_country.empty:
            plexos_results = None
            print('No data extracted for', country)
        else:
            plexos_results = llm_call_with_json(default_data_prompt, default_data_context, 
                                  model=self.config.sota_model)
        
        return plexos_results
    
    def _build_data_analysis_context(self, instruction: str, json_data: str, 
                                    chapter: str) -> str:
        """Build context for data analysis."""
        return f"""{self.config.main_context}. 
                You are currently drafting chapter {chapter}.
                Instruction: {instruction}.
                Do not end the chapter with any conclusion, unless stated.
                Do not include a header in the response.
                Use neutral language, do not overuse adjectives. Use simple language.
                """


class ReportStructureManager:
    """Manages the report structure and data organization with thread-safety."""
    
    def __init__(self, config: ProjectConfig):
        self.config = config
        self.report_structure: Dict = {}
        self._lock = Lock()  # Thread-safe access to report_structure
    
    def structure_dictionary(self, prompt_sheet: Dict, objective_id: str, task_id: str,
                           sub_task_id: str, objective_title: str, task_title: str,
                           sub_task_title: str, main_copy: str, summary: Optional[str],
                           country: str, refinement_service: TextRefinementService) -> Dict:
        """Structure report data into hierarchical dictionary (thread-safe)."""
        
        # Perform refinement and feedback outside the lock (these are I/O intensive)
        if self.config.add_final_polish:
            final_copy = refinement_service.refine_and_proofread(sub_task_id, main_copy, 
                                                                 prompt_sheet, country, 'main')
        else:
            final_copy = main_copy

        if self.config.add_feedback and country != 'EU':
            report_extract_dict = {
                'objective_title': objective_title,
                'task_title': task_title,
                'sub_task_title': sub_task_title,
                'country': country,
                'final_copy': final_copy
            }
            feedback = refinement_service.feedback_sense_check(
                report_extract_dict, 
                self.config.main_context, 
                main_copy
            )
            print('Feedback for', sub_task_id, 'for', country, feedback)
        else:
            feedback = None
        
        # Only lock when modifying the shared data structure
        with self._lock:
            if objective_id not in self.report_structure:
                self.report_structure[objective_id] = {
                    'Objective': objective_id,
                    'tasks': {}
                }

            if task_id not in self.report_structure[objective_id]['tasks']:
                self.report_structure[objective_id]['tasks'][task_id] = {
                    'task_id': task_id,
                    'task_title': task_title,
                    'sub_tasks': {}
                }

            if sub_task_id not in self.report_structure[objective_id]['tasks'][task_id]['sub_tasks']:
                self.report_structure[objective_id]['tasks'][task_id]['sub_tasks'][sub_task_id] = {
                    'sub_task_id': sub_task_id,
                    'sub_task_title': sub_task_title,
                    'country': {}
                }

            self.report_structure[objective_id]['tasks'][task_id]['sub_tasks'][sub_task_id]['country'][country] = {
                'final_copy': final_copy,
                'objective_title': objective_title,
                'task_title': task_title,
                'sub_task_title': sub_task_title,
                'summary': summary,
                'feedback': feedback
            }
        
        return self.report_structure
    
    def export_to_dataframe(self, objective_id: str, file_manager: FileManager) -> pd.DataFrame:
        """Export report structure to DataFrame and save to CSV (thread-safe)."""
        result_list = []
        
        # Read from shared structure with lock
        with self._lock:
            structure_copy = dict(self.report_structure)  # Create a copy
        
        for obj_id, objective_data in structure_copy.items():
            for task_id, task_data in objective_data['tasks'].items():
                for sub_task_id, sub_task_data in task_data['sub_tasks'].items():
                    for country, country_data in sub_task_data['country'].items():
                        result_list.append({
                            'objective_id': obj_id,
                            'task_id': task_id,
                            'sub_task_id': sub_task_id,
                            'country': country,
                            'final_copy': country_data['final_copy'],
                            'objective_title': country_data['objective_title'],
                            'task_title': country_data['task_title'],
                            'sub_task_title': country_data['sub_task_title'],
                            'summary': country_data['summary']
                        })

        if result_list:
            result_df = pd.DataFrame(result_list)
            result_df = result_df[result_df['objective_id'] == objective_id]
        else:
            result_df = pd.DataFrame(columns=[
                'objective_id', 'task_id', 'sub_task_id', 'country', 'final_copy',
                'objective_title', 'task_title', 'sub_task_title', 'summary'
            ])

        # Merge with existing data
        file_path = os.path.join(self.config.report_output_location, 
                                self.config.project_name, 
                                f'objective_{objective_id}_results.csv')
        
        if os.path.exists(file_path):
            existing_df = pd.read_csv(file_path)
            missing_rows = existing_df[~existing_df['sub_task_id'].isin(result_df['sub_task_id'])]
            result_df = pd.concat([result_df, missing_rows], ignore_index=True)

        file_manager.save_file(file_path, result_df, file_type='csv')
        return result_df
    
    def rebuild_report(self, file_manager: FileManager) -> pd.DataFrame:
        """Rebuild complete report from saved objective files."""
        report_structure = pd.DataFrame()
        
        for objective in range(1, 6):
            file_path = os.path.join(self.config.report_output_location,
                                    self.config.project_name,
                                    f'objective_{objective}_results.csv')
            if os.path.exists(file_path):
                objective_df = pd.read_csv(file_path)
                report_structure = pd.concat([report_structure, objective_df])
        
        return report_structure


class ReportDrafter:
    """Handles the main report drafting operations."""
    
    def __init__(self, config: ProjectConfig, 
                 data_analysis_service: DataAnalysisService,
                 refinement_service: TextRefinementService,
                 structure_manager: ReportStructureManager):
        self.config = config
        self.data_analysis = data_analysis_service
        self.refinement = refinement_service
        self.structure_manager = structure_manager
    
    def add_data_to_report(self, prompt_sheet: Dict, task_header: str, 
                          sub_task_id: str, objective_id: str, objective_title: str,
                          task_title: str, sub_task_title: str, standard_guidelines: str,
                          research_notes: str, detailed_context: Any, 
                          additional_context: Any, sub_task_description: str,
                          location: str, task_id: str) -> Dict:
        """Add analyzed data to report structure."""
        country_context = f"""
                        {self.config.main_context}.
                        You are drafting section for objective: {objective_title}, 
                        task {task_title}, sub-task {sub_task_title}.
                        Consider your position in the document to determine how to write the section.
                    """
        
        country_prompt = f"""
                    Write the report chapter {sub_task_id}: {task_header}, {sub_task_title} 
                    on the data analysis based on the information provided.
                    Write a report for the country: {location}.
                    Instruction: {sub_task_description}.
                    Section guidelines: {standard_guidelines}.
                    Research notes: {research_notes}.
                    Context: {detailed_context}.
                    Data extracts: {additional_context}.   
                    Illustrate with numbers if provided.   
                    Make it comprehensive, think about connections in the data.
                    Don't speak about policy or recommendations, just data and analysis.
                    Write in prose, no additional formatting.
                    Do not include a header in the response.
                    Only consider country ranking if significant (top 5 or bottom 5).
                """
        
        first_draft = llm_call_with_json(country_prompt, country_context, model=self.config.writer_model)
        first_draft = first_draft.replace('—', ', ')

        if self.config.add_summary:
            summary_context = f"""
                                {self.config.main_context}.
                                You are asked to summarize the first draft.
                            """
            summary_prompt = f"""
                            Please summarize the first draft in a concise manner.
                            Here is the first draft: {first_draft}
                            """
            summary = llm_call_with_json(summary_prompt, summary_context, model=self.config.base_model)
        else:
            summary = None

        try:
            country = self.config.countries[location]
        except KeyError:
            country = location

        report_structure = self.structure_manager.structure_dictionary(
            prompt_sheet, objective_id, task_id, sub_task_id, objective_title,
            task_title, sub_task_title, first_draft, summary, country, 
            self.refinement
        )
        
        return report_structure


class AnalysisOrchestrator:
    """Orchestrates the analysis of default and additional data."""
    
    def __init__(self, config: ProjectConfig,
                 plexos_extractor: PlexosDataExtractor,
                 etm_extractor: ETMDataExtractor,
                 data_analysis_service: DataAnalysisService,
                 search_service: InternetSearchService):
        self.config = config
        self.plexos_extractor = plexos_extractor
        self.etm_extractor = etm_extractor
        self.data_analysis = data_analysis_service
        self.search_service = search_service
    
    def analyse_default_data(self, prompt_sheet: Dict, sub_task_id: str, 
                            detailed_context: Dict) -> Dict:
        """Analyze default chart data."""
        row_extract = prompt_sheet['Default Charts'].set_index('sub_task_id', 
                                                               drop=False).loc[sub_task_id]
        description = row_extract['Description']
        chart_type = row_extract['Image']

        if chart_type == 'plexos_chart':
            plexos_extract = self.plexos_extractor.extract(row_extract, 'Default Charts', 
                                                           sub_task_id)
            plexos_extract = plexos_extract[plexos_extract['value'] != 0]
            
            for country in self.config.countries:
                try:
                    plexos_results = self.data_analysis.process_plexos_data(
                        country, plexos_extract, description, prompt_sheet, row_extract
                    )
                    detailed_context[country] = {'Default Data - PLEXOS': plexos_results}
                except Exception as e:
                    print(f'Error in analysing default data for {sub_task_id}', e, country)

        elif chart_type == 'ETM':
            for country in self.config.countries:
                print('Analyzing default ETM data for', country)
                default_charts = row_extract
                etm_extract_country = self.etm_extractor.extract(
                    sub_task_id, detailed_context, default_charts, country, description
                )
                context = self.config.main_context
                prompt = f"""Please give a concise summary of the data extracted from the ETM database.
                        Here are the Report extracts: {etm_extract_country}. return only the copy
                        """
                etm_result_analysis = llm_call_with_json(prompt, context, model=self.config.base_model)
                detailed_context[country] = {'Default Data - ETM': etm_result_analysis}

        return detailed_context
    
    def analyse_additional_data(self, prompt_sheet: Dict, sub_task_id: str, 
                               detailed_context: Dict) -> Dict:
        """Analyze additional external data sources."""
        print('Analyzing additional data for', sub_task_id)
        guidelines = prompt_sheet['External Search'].loc[
            prompt_sheet['External Search'].index == sub_task_id
        ]
        guidelines.set_index('Unique_ID', inplace=True)

        country_results = {}
        level_results = {}
        title_results = {}

        all_additional_information_str = ' '.join(str(guidelines['additional information']))

        for guideline_id in guidelines.index:
            current_guideline = guidelines.loc[guideline_id]
            data_source = current_guideline['Data Source']
            level = current_guideline['Level']
            prompt = current_guideline['Prompt']
            description = current_guideline['Description']
            title = current_guideline['Title']

            if level == 'country':
                for country in self.config.countries:
                    if country not in country_results:
                        country_results[country] = []

                    if data_source == 'Internet Search':
                        print(f'Performing Internet Search for task: {sub_task_id} for country: {country}')
                        additional_information = current_guideline['additional information']
                        additional_information = f'{additional_information} for {self.config.countries[country]}'
                        search_result = self.search_service.search(
                            country, additional_information, prompt, description
                        )
                        country_results[country].append(search_result)

                    elif data_source == 'ETM':
                        print(f'Extracting ETM data for {country}')
                        etm_results = self.etm_extractor.extract(
                            sub_task_id, detailed_context, current_guideline, country, prompt
                        )
                        country_results[country].append(etm_results)
                    
                    elif data_source == 'PLEXOS':
                        plexos_extract = self.plexos_extractor.extract(
                            current_guideline, 'External Search', sub_task_id
                        )
                        plexos_results = self.data_analysis.process_plexos_data(
                            country, plexos_extract, description, prompt_sheet, current_guideline
                        )
                        country_results[country].append(plexos_results)

            else:
                if data_source == 'Internet Search':
                    print('Analysing additional data for EU')
                    search_result = self.search_service.search(
                        'EU', current_guideline['additional information'], prompt, description
                    )
                    level_results[level] = level_results.get(level, []) + [search_result]

                elif data_source == 'LLM Call':
                    print('Analysing LLM data')
                    llm_results = llm_call_with_json(title, description, model=self.config.base_model)
                    title_results[title] = title_results.get(title, []) + [llm_results]

        # Synthesize results
        for country in country_results:
            if country_results[country]:
                all_text = ' '.join(str(result) for result in country_results[country])
                synthesis_prompt = f"""
                                    Research summary for {country}.
                                    Title: {title}.
                                    Description: {description}.
                                    Prompt: {prompt}.
                                    Additional research: {all_additional_information_str}.
                                    Information obtained: {all_text}
                                    """
                country_synthesis = llm_call_with_json(synthesis_prompt, self.config.main_context, 
                                        model=self.config.base_model)
                detailed_context[country] = {'synthesis': country_synthesis}

        return detailed_context


class ReportGenerationPipeline:
    """Main pipeline for report generation orchestration."""
    
    def __init__(self, config: ProjectConfig, prompt_sheet: Dict, 
                 etm_data: pd.DataFrame, nodalsplit: pd.DataFrame):
        self.config = config
        self.prompt_sheet = prompt_sheet
        
        # Initialize services
        self.file_manager = FileManager()
        self.plexos_extractor = PlexosDataExtractor(config)
        self.etm_extractor = ETMDataExtractor(config, etm_data)
        self.search_service = InternetSearchService(config)
        self.data_analysis = DataAnalysisService(config, self.plexos_extractor, nodalsplit)
        self.refinement_service = TextRefinementService(config)
        self.structure_manager = ReportStructureManager(config)
        self.report_drafter = ReportDrafter(config, self.data_analysis, 
                                           self.refinement_service, self.structure_manager)
        self.analysis_orchestrator = AnalysisOrchestrator(
            config, self.plexos_extractor, self.etm_extractor, 
            self.data_analysis, self.search_service
        )
        self.exporter = ReportExporter(self.file_manager)
    
    def process_sub_tasks(self, sub_tasks: pd.DataFrame, objective_id: str, 
                         objective_title: str, task_title: str, task_id: str) -> Dict:
        """Process all sub-tasks for a given task."""
        for sub_task_id in sub_tasks.index:
            if not self.config.sub_task_list or sub_task_id in self.config.sub_task_list:
                sub_task_title = sub_tasks.loc[sub_task_id, 'Title']
                sub_task_description = sub_tasks.loc[sub_task_id, 'Description']
                geographical_level = sub_tasks.loc[sub_task_id, 'Geographic Level']
                task_header = sub_tasks.loc[sub_task_id, 'Task Header']
                
                print(sub_task_id, sub_task_description)

                detailed_context = {}
                additional_context = {}
                
                standard_guidelines = self.prompt_sheet['Text Guidelines'].loc[
                    sub_task_id, 'Standard'
                ]
                research_notes = self.prompt_sheet['Text Guidelines'].loc[
                    sub_task_id, 'Research Notes'
                ]

                # Extract default data
                for x in self.prompt_sheet['Default Charts']['sub_task_id']:
                    if x == sub_task_id:
                        try:
                            detailed_context = self.analysis_orchestrator.analyse_default_data(
                                self.prompt_sheet, sub_task_id, detailed_context
                            )
                        except Exception as e:
                            print(f'Error in analysing default data for {sub_task_id}', e)

                # Extract additional data
                if sub_task_id in self.prompt_sheet['External Search'].index.values:
                    try:
                        additional_context = self.analysis_orchestrator.analyse_additional_data(
                            self.prompt_sheet, sub_task_id, additional_context
                        )
                    except Exception as e:
                        print(f'Error in analysing additional data for {sub_task_id}', e)

                # Draft report sections
                if geographical_level == 'country':
                    countries = [
                        country for country in detailed_context
                        if detailed_context[country].get('Default Data - PLEXOS') is not None
                        or detailed_context[country].get('Default Data - ETM') is not None
                    ]
                    
                    if detailed_context:
                        for country in countries:
                            print('Drafting section', sub_task_id, sub_task_title, 'for', country)
                            try:
                                detailed_context_country = detailed_context[country]
                                additional_context_country = additional_context.get(country)
                                
                                self.structure_manager.report_structure = self.report_drafter.add_data_to_report(
                                    self.prompt_sheet, task_header, sub_task_id, 
                                    objective_id, objective_title, task_title, sub_task_title,
                                    standard_guidelines, research_notes, detailed_context_country,
                                    additional_context_country, sub_task_description, 
                                    country, task_id
                                )
                            except Exception as e:
                                print(f'Error in Drafting section {sub_task_id}', e, country)
                else:
                    print('Drafting section', sub_task_id, sub_task_title)
                    self.structure_manager.report_structure = self.report_drafter.add_data_to_report(
                        self.prompt_sheet, task_header, sub_task_id, 
                        objective_id, objective_title, task_title, sub_task_title,
                        standard_guidelines, research_notes, detailed_context,
                        additional_context, sub_task_description, geographical_level, task_id
                    )
            else:
                print('Skipping sub-task', sub_task_id)
        
        return self.structure_manager.report_structure
    
    def _process_single_task(self, task_id: str, objective_id: str, 
                            objective_title: str, context: Dict) -> None:
        """Process a single task (for concurrent execution)."""
        task_title = self.prompt_sheet['Tasks'].loc[task_id, 'Title']
        task_description = self.prompt_sheet['Tasks'].loc[task_id, 'Description']
        task_context = context.copy()
        task_context['Task'] = task_description
        
        print(f"Processing task {task_id}: {task_description}")
        
        sub_tasks = self.prompt_sheet['Sub Tasks'][
            self.prompt_sheet['Sub Tasks']['Task_section_id'] == task_id
        ]
        
        self.structure_manager.report_structure = self.process_sub_tasks(
            sub_tasks, objective_id, objective_title, task_title, task_id
        )
    
    def parse_objectives(self, objective_id: str, context: Dict) -> Dict:
        """Parse and process objectives."""
        objective_title = self.prompt_sheet['Objectives'].loc[objective_id, 'Title']
        objective_description = self.prompt_sheet['Objectives'].loc[objective_id, 'Description']
        
        print(objective_id, objective_description)
        context['Objective'] = objective_description
        
        tasks = self.prompt_sheet['Tasks'][
            self.prompt_sheet['Tasks']['Objective ID'] == objective_id
        ]
        
        valid_task_ids = [task_id for task_id in tasks.index if task_id > 0]

        if self.config.run_task_concurrent and len(valid_task_ids) > 1:
            print(f"Running {len(valid_task_ids)} tasks concurrently for objective {objective_id}")
            with concurrent.futures.ThreadPoolExecutor(max_workers=len(valid_task_ids)) as executor:
                futures = []
                for task_id in valid_task_ids:
                    future = executor.submit(
                        self._process_single_task,
                        task_id, objective_id, objective_title, context
                    )
                    futures.append(future)
                
                # Wait for all tasks to complete
                concurrent.futures.wait(futures)
                
                # Check for exceptions
                for future in futures:
                    try:
                        future.result()
                    except Exception as e:
                        print(f"Error in task processing: {e}")
        else:
            # Sequential processing
            for task_id in valid_task_ids:
                task_title = self.prompt_sheet['Tasks'].loc[task_id, 'Title']
                task_description = self.prompt_sheet['Tasks'].loc[task_id, 'Description']
                context['Task'] = task_description
                
                print(task_id, task_description)
                
                sub_tasks = self.prompt_sheet['Sub Tasks'][
                    self.prompt_sheet['Sub Tasks']['Task_section_id'] == task_id
                ]
                
                self.structure_manager.report_structure = self.process_sub_tasks(
                    sub_tasks, objective_id, objective_title, task_title, task_id
                )
        
        self.structure_manager.export_to_dataframe(objective_id, self.file_manager)
        return self.structure_manager.report_structure
    
    def conclusion_next_steps(self) -> None:
        """Generate conclusion and next steps sections."""
        print('Generating conclusion and next steps')
        # Implementation similar to original but using class methods
        # This is a simplified version - full implementation would follow original logic
        pass
    
    def _process_single_objective(self, objective_id: str) -> None:
        """Process a single objective (for concurrent execution)."""
        context = {}
        print(f'Processing objective {objective_id}')
        self.parse_objectives(objective_id, context)
    
    def run(self) -> None:
        """Execute the complete report generation pipeline."""
        start_time = time.time()
        start_time_str = time.strftime('%H:%M:%S', time.localtime())
        print(f"Report generation started at {start_time_str}")
        
        warnings.filterwarnings('ignore')
        
        # Ensure output directory exists
        output_dir = os.path.join(self.config.report_output_location, 
                                 self.config.project_name)
        self.file_manager.ensure_directory_exists(output_dir)
        
        # Main report generation
        if self.config.run_main_report:
            objectives = self.prompt_sheet['Objectives'].index.unique()
            
            # Filter objectives based on config
            valid_objectives = [
                obj_id for obj_id in objectives 
                if obj_id in self.config.objective_list or not self.config.objective_list
            ]
            
            if self.config.run_objectives_concurrent and len(valid_objectives) > 1:
                print(f"Running {len(valid_objectives)} objectives (chapters) concurrently")
                with concurrent.futures.ThreadPoolExecutor(max_workers=len(valid_objectives)) as executor:
                    futures = []
                    for objective_id in valid_objectives:
                        future = executor.submit(self._process_single_objective, objective_id)
                        futures.append(future)
                    
                    # Wait for all objectives to complete
                    concurrent.futures.wait(futures)
                    
                    # Check for exceptions
                    for future in futures:
                        try:
                            future.result()
                        except Exception as e:
                            print(f"Error in objective processing: {e}")
            else:
                # Sequential processing
                context = {}
                for objective_id in valid_objectives:
                    print(f'Processing objective {objective_id}')
                    self.parse_objectives(objective_id, context)
        
        # Generate conclusions
        if self.config.conclude_report:
            self.conclusion_next_steps()
        
        # Build final report
        final_report = self.structure_manager.rebuild_report(self.file_manager)
        if 'summary' in final_report.columns:
            final_report.drop(columns=['summary'], inplace=True)
        
        # Save final report
        file_path = os.path.join(self.config.report_output_location,
                                self.config.project_name,
                                f'{self.config.project_name}.csv')
        self.file_manager.save_file(file_path, final_report, file_type='csv')
        
        # Create Word document
        word_path = os.path.join(self.config.report_output_location, 
                                'Joule_Report_V2.docx')
        self.exporter.create_word_file(final_report, output_path=word_path)
        
        # Print timing
        end_time = time.time()
        end_time_str = time.strftime('%H:%M:%S', time.localtime())
        elapsed_hours = (end_time - start_time) / 3600
        print(f'Time taken: {elapsed_hours:.2f} hours')
        print(f'Start time: {start_time_str}, End time: {end_time_str}')


def load_configuration(project: str = 'Joule') -> Tuple[ProjectConfig, pd.DataFrame, pd.DataFrame]:
    """Load configuration and data files."""
    # File paths
    if project == 'Joule':
        file_path = r'C:\Users\ENTSOE\Tera-joule\Terajoule - Terajoule\Projects\Sectoral Model\Website\Joule_prompt_sheet_.xlsx'
    elif project == 'Medtso':
        file_path = r'C:\Users\ENTSOE\Tera-joule\Terajoule - Terajoule\Projects\Sectoral Model\Website\TEASIMED_IoSN_prompt_sheet.xlsx'
    else:
        raise ValueError(f"Unknown project: {project}")
    
    # Load data
    etm = pd.read_csv(r'external_resources\model_databases\joule_model\etm_2024\etm_extract_joule_report.csv')
    nodalsplit = pd.read_excel(
        r'src\EMIL\demand\Input\Population and Industrial Sizes.xlsx', 
        sheet_name='ehighway_nut2 conversion'
    )
    
    # Load country borders
    country_borders_file = r'config\topology\country_borders.yaml'
    with open(country_borders_file, 'r') as f:
        country_borders = yaml.safe_load(f) or {}
    
    # Load AI models config
    default_ai_models_file = r'config\default_ai_models.yaml'
    with open(default_ai_models_file, 'r') as f:
        ai_models_config = yaml.safe_load(f)
    
    # Load prompt sheet
    prompt_sheet = eps.import_excel_sheets_as_dict(file_path)
    
    # Countries dictionary
    countries = {
        "AT": "Austria", "BE": "Belgium", "BG": "Bulgaria", "CY": "Cyprus", 
        "CZ": "Czech Republic", "DE": "Germany", "DK": "Denmark", "EE": "Estonia", 
        "ES": "Spain", "FI": "Finland", "FR": "France", "GR": "Greece", 
        "HR": "Croatia", "HU": "Hungary", "IE": "Ireland", "IT": "Italy", 
        "LT": "Lithuania", "LU": "Luxembourg", "LV": "Latvia", "MT": "Malta", 
        "NL": "Netherlands", "PL": "Poland", "PT": "Portugal", "RO": "Romania", 
        "SE": "Sweden", "SI": "Slovenia", "SK": "Slovakia", "UK": "United Kingdom"
    }
    
    # Tone and feel
    tone_and_feel = {
        'purpose': 'Educate',
        'target_audience': 'Policy Makers, TSO, Energy Project Owner, NGOs',
        'tone_voice': 'Slightly pro nuclear',
        'format': 'Energy Report',
    }
    
    # Create config
    config = ProjectConfig(
        project_name=str(prompt_sheet['Read me']['Project Name'][0]),
        file_path=file_path,
        base_location=r'external_resources\model_databases\joule_model',
        report_output_location=fr'functions\copywriting_functions\Generated Reports\{project}',
        model_name='TJ Dispatch_Future_Nuclear+',
        tone_and_feel=tone_and_feel,
        countries=countries,
        main_context=str(prompt_sheet['Read me']['Context']),
        base_model=ai_models_config.get("base_model", "gpt-5-mini"),
        pro_model=ai_models_config.get("pro_model", "gpt-5"),
        sota_model=ai_models_config.get("sota_model", "gpt-5"),
        writer_model=ai_models_config.get("writer_model", "gpt-5"),
        country_borders=country_borders
    )
    
    # Prepare nodalsplit
    nodalsplit['Country'] = nodalsplit['Node'].str[0:2]
    nodalsplit.drop(columns=['Position'], inplace=True, errors='ignore')
    
    return config, etm, nodalsplit


def main():
    """Main entry point for the report generation system."""
    # Load configuration
    config, etm_data, nodalsplit = load_configuration(project='Joule')
    
    # Load prompt sheet
    prompt_sheet = eps.import_excel_sheets_as_dict(config.file_path)
    
    # Create and run pipeline
    pipeline = ReportGenerationPipeline(config, prompt_sheet, etm_data, nodalsplit)
    pipeline.run()


if __name__ == "__main__":
    main()
